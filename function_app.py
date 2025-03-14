import azure.functions as func
import logging
import json
from azure.storage.blob import BlobServiceClient
from azure.data.tables import TableServiceClient, TableClient
import os
import tempfile
import PyPDF2
import docx
import uuid
import pandas as pd
import openpyxl
import requests
import time
from io import BytesIO
from datetime import datetime

app = func.FunctionApp(http_auth_level=func.AuthLevel.FUNCTION)

@app.route(route="ProcessResume", methods=["POST"])
def ProcessResume(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('レジュメアップロード処理が開始されました')
    
    try:
        # Blobストレージの接続設定
        connection_string = os.environ["StorageConnectionString"]
        container_name = "resumes"
        
        # POSTリクエストからファイルを取得
        file_data = req.get_body()
        content_type = req.headers.get('content-type', '')
        
        if not file_data:
            return func.HttpResponse(
                json.dumps({"error": "ファイルが見つかりません"}),
                mimetype="application/json",
                status_code=400
            )
        
        # ファイル形式を判定
        file_extension = ""
        content_type = req.headers.get('content-type', '').lower()

        if 'pdf' in content_type:
            file_extension = ".pdf"
        elif ('spreadsheetml' in content_type or 'excel' in content_type or
              'sheet' in content_type or 'xlsx' in content_type or 'xls' in content_type):
            if 'xlsx' in content_type:
                file_extension = ".xlsx"
            else:
                file_extension = ".xls"
        elif ('wordprocessingml' in content_type or 'document' in content_type or
              'docx' in content_type or 'doc' in content_type):
            if 'docx' in content_type:
                file_extension = ".docx"
            else:
                file_extension = ".doc"
        else:
            filename = req.headers.get('x-ms-file-name', '')
            if filename.lower().endswith('.pdf'):
                file_extension = ".pdf"
            elif filename.lower().endswith('.docx'):
                file_extension = ".docx"
            elif filename.lower().endswith('.xlsx'):
                file_extension = ".xlsx"
            elif filename.lower().endswith('.xls'):
                file_extension = ".xls"
            else:
                return func.HttpResponse(
                    json.dumps({"error": f"サポートされていないファイル形式です: {content_type}"}),
                    mimetype="application/json",
                    status_code=400
                )
        
        # 一意のファイル名を生成
        file_id = str(uuid.uuid4())
        file_name = file_id + file_extension
        
        # Blobストレージに保存
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name)
        
        if not container_client.exists():
            container_client.create_container()
        
        blob_client = container_client.get_blob_client(file_name)
        blob_client.upload_blob(file_data, overwrite=True)
        
        # テキスト抽出処理
        text_content = extract_text(file_data, file_extension)
        
        return func.HttpResponse(
            json.dumps({
                "fileId": file_id,
                "fileName": file_name,
                "textContent": text_content,
                "status": "success"
            }),
            mimetype="application/json",
            status_code=200
        )
        
    except Exception as e:
        logging.error(f"エラー発生: {str(e)}")
        return func.HttpResponse(
            json.dumps({"error": str(e)}),
            mimetype="application/json",
            status_code=500
        )

def extract_text(file_data, file_extension):
    """ファイルからテキストを抽出する関数"""
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=file_extension)
        temp_file.write(file_data)
        temp_file.close()
        
        text_content = ""
        
        if file_extension == ".pdf":
            with open(temp_file.name, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text_content += (page.extract_text() or "") + "\n"
        
        elif file_extension == ".docx":
            doc = docx.Document(temp_file.name)
            for para in doc.paragraphs:
                text_content += para.text + "\n"
        
        elif file_extension in [".xlsx", ".xls"]:
            all_sheets_text = []
            xls = pd.ExcelFile(temp_file.name)
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(temp_file.name, sheet_name=sheet_name)
                all_sheets_text.append(f"===== シート: {sheet_name} =====")
                sheet_text = df.to_string(index=False)
                all_sheets_text.append(sheet_text)
                all_sheets_text.append("\n")
            text_content = "\n".join(all_sheets_text)
            
            if file_extension == ".xlsx":
                try:
                    workbook = openpyxl.load_workbook(temp_file.name, data_only=True)
                    for sheet_name in workbook.sheetnames:
                        sheet = workbook[sheet_name]
                        for row in sheet.iter_rows():
                            for cell in row:
                                if cell.comment:
                                    text_content += f"\nコメント ({sheet_name} {cell.coordinate}): {cell.comment.text}"
                except Exception as e:
                    logging.warning(f"Excel詳細情報抽出エラー: {str(e)}")
        
        os.unlink(temp_file.name)
        return text_content
    
    except Exception as e:
        logging.error(f"テキスト抽出エラー: {str(e)}")
        raise e

@app.route(route="AnalyzeWithGPT", methods=["POST"])
def AnalyzeWithGPT(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('GPT API連携関数が呼び出されました')
    
    # リクエストヘッダーをログに記録
    request_headers = dict(req.headers)
    logging.info(f"リクエストヘッダー: {request_headers}")
    
    # リクエストボディを取得（生のバイトデータ）
    request_body_bytes = req.get_body()
    
    # Content-Typeを確認
    content_type = req.headers.get('content-type', '').lower()
    logging.info(f"Content-Type: {content_type}")
    
    # 文字列に変換してログに記録
    body_text = request_body_bytes.decode('utf-8', errors='replace')
    body_preview = body_text[:1000] + ('...' if len(body_text) > 1000 else '')
    logging.info(f"受信したリクエストボディ (デコード): {body_preview}")
    
    try:
        # JSONとして解析
        if 'application/json' in content_type:
            try:
                req_body = req.get_json()
                logging.info(f"JSONとして解析されたリクエストボディ (キー): {list(req_body.keys())}")
            except ValueError as e:
                # 手動でJSONを解析してみる
                try:
                    req_body = json.loads(body_text)
                    logging.info("手動JSON解析成功")
                except json.JSONDecodeError as json_err:
                    logging.error(f"JSON解析エラー: {str(json_err)}")
                    return func.HttpResponse(
                        json.dumps({
                            "error": f"JSONの解析に失敗しました: {str(json_err)}",
                            "request_preview": body_preview[:200]
                        }),
                        mimetype="application/json",
                        status_code=400
                    )
        else:
            logging.error(f"サポートされていないContent-Type: {content_type}")
            return func.HttpResponse(
                json.dumps({"error": f"サポートされていないContent-Type: {content_type}"}),
                mimetype="application/json",
                status_code=400
            )
        
        # データの検証
        resume_text = req_body.get('resumeText')
        jd_data = req_body.get('jdData', {})
        
        # 入力データの検証
        if not resume_text:
            logging.error("レジュメテキストが提供されていません")
            return func.HttpResponse(
                json.dumps({"error": "レジュメテキストが提供されていません"}),
                mimetype="application/json",
                status_code=400
            )
        
        if not isinstance(jd_data, dict):
            logging.error(f"不正なJDデータ形式: {type(jd_data)}")
            return func.HttpResponse(
                json.dumps({"error": "JDデータは辞書形式である必要があります"}),
                mimetype="application/json",
                status_code=400
            )
        
        logging.info("GPT API呼び出し開始")
        
        # GPT API呼び出し（リトライロジック付き）
        max_retries = 3
        retry_count = 0
        last_error = None
        
        while retry_count < max_retries:
            try:
                analysis_result = call_gpt_api(resume_text, jd_data)
                
                # エラーがある場合は再試行
                if "error" in analysis_result:
                    error_msg = analysis_result["error"]
                    if "レート制限" in error_msg or "rate limit" in error_msg.lower():
                        retry_count += 1
                        wait_time = min(2 ** retry_count, 30)
                        logging.warning(f"GPT APIレート制限エラー。{wait_time}秒後に再試行 ({retry_count}/{max_retries})")
                        time.sleep(wait_time)
                        continue
                
                # エラーがなければ、または再試行対象外のエラーならば結果を返す
                break
                
            except Exception as e:
                last_error = e
                retry_count += 1
                wait_time = min(2 ** retry_count, 30)
                logging.error(f"GPT API呼び出しエラー: {str(e)}. {wait_time}秒後に再試行 ({retry_count}/{max_retries})")
                time.sleep(wait_time)
        
        # すべてのリトライが失敗した場合
        if retry_count == max_retries and last_error:
            logging.error(f"すべてのリトライが失敗しました: {str(last_error)}")
            return func.HttpResponse(
                json.dumps({"error": f"GPT APIへの接続に失敗しました: {str(last_error)}"}),
                mimetype="application/json",
                status_code=500
            )
        
        logging.info("GPT API呼び出し完了")
        
        # CORS対応ヘッダーを追加
        headers = {
            "Access-Control-Allow-Origin": "*",
            "Access-Control-Allow-Methods": "POST, OPTIONS",
            "Access-Control-Allow-Headers": "Content-Type, Authorization"
        }
        
        return func.HttpResponse(
            json.dumps(analysis_result, ensure_ascii=False),
            mimetype="application/json",
            headers=headers,
            status_code=200
        )
    except Exception as e:
        logging.error(f"予期せぬエラー: {str(e)}")
        return func.HttpResponse(
            json.dumps({"error": str(e)}),
            mimetype="application/json",
            status_code=500
        )

def call_gpt_api(resume_text, jd_data):
    """GPT APIを呼び出してレジュメを分析する関数"""
    try:
        # APIキーの取得 - 複数の環境変数名をチェック
        api_key = os.environ.get("OPENAI_API_KEY", os.environ.get("OPENAI_HERE", ""))
        if not api_key:
            logging.error("環境変数 OPENAI_API_KEY または OPENAI_HERE が設定されていません")
            return {"error": "APIキーが設定されていません"}
            
        api_url = "https://api.openai.com/v1/chat/completions"
        
        # テキストの長さを制限
        max_tokens = 8000  # より安全な長さに調整
        if len(resume_text) > max_tokens:
            resume_text = resume_text[:max_tokens] + "...(省略)"
            
        prompt = create_analysis_prompt(resume_text, jd_data)
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # レスポンスフォーマットの検証
        use_json_format = True
        
        try:
            # テスト用に小さなリクエストを送信して、response_formatがサポートされているか確認
            test_payload = {
                "model": "gpt-4",
                "messages": [{"role": "user", "content": "Say hello in JSON"}],
                "max_tokens": 20,
                "response_format": {"type": "json_object"}
            }
            test_response = requests.post(api_url, headers=headers, json=test_payload, timeout=10)
            
            if test_response.status_code == 400 and "response_format" in test_response.text:
                use_json_format = False
                logging.warning("API doesn't support response_format, disabling...")
        except Exception as e:
            use_json_format = False
            logging.warning(f"API テスト中にエラーが発生しました、response_format を無効化: {str(e)}")
        
        # メインのペイロード
        payload = {
            "model": "gpt-4",
            "messages": [
                {"role": "system", "content": "あなたは履歴書を分析する専門家です。必ず有効なJSONのみを出力し、JSON以外のテキストは出力しないでください。"},
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 2048,  # より安全な値に調整
            "temperature": 0.2
        }
        
        # 最新のAPIでサポートされている場合のみ追加
        if use_json_format:
            payload["response_format"] = {"type": "json_object"}
        
        logging.info("OpenAI APIリクエスト準備完了")
        logging.info(f"使用モデル: {payload['model']}, response_format 使用: {use_json_format}")
        
        # リクエスト送信（タイムアウト設定）
        response = requests.post(api_url, headers=headers, json=payload, timeout=180)
        
        # エラー処理
        if response.status_code != 200:
            error_detail = response.text[:500] if response.text else "詳細なし"
            logging.error(f"OpenAI API応答エラー: Status={response.status_code}, Response={error_detail}")
            
            if response.status_code == 401:
                return {"error": "OpenAI APIの認証に失敗しました。APIキーを確認してください。"}
            elif response.status_code == 429:
                return {"error": "OpenAI APIのレート制限に達しました。しばらく時間をおいて再試行してください。"}
            elif response.status_code == 500:
                return {"error": "OpenAI APIでサーバーエラーが発生しました。しばらく時間をおいて再試行してください。"}
            else:
                return {"error": f"OpenAI APIエラー: {response.status_code} - {error_detail}"}
            
        # レスポンスの解析
        result = response.json()
        analysis_text = result['choices'][0]['message']['content']
        
        logging.info(f"GPT応答テキスト長: {len(analysis_text)} 文字")
        logging.info(f"GPT応答テキストプレビュー: {analysis_text[:100]}...")
        
        # JSONデータの抽出と解析
        try:
            # まず単純にJSONとして解析を試みる
            analysis_data = json.loads(analysis_text)
            logging.info("JSON直接解析成功")
        except json.JSONDecodeError as json_err:
            logging.warning(f"JSON直接解析エラー: {str(json_err)}")
            
            # JSON部分を抽出して解析を試みる
            try:
                json_start = analysis_text.find('{')
                json_end = analysis_text.rfind('}') + 1
                
                if json_start >= 0 and json_end > 0:
                    json_str = analysis_text[json_start:json_end]
                    logging.info(f"抽出されたJSON文字列: {json_str[:100]}...")
                    analysis_data = json.loads(json_str)
                    logging.info("JSON部分解析成功")
                else:
                    logging.error("JSONが見つかりません")
                    analysis_data = {"raw_analysis": analysis_text}
            except json.JSONDecodeError as extract_err:
                logging.error(f"JSON部分解析エラー: {str(extract_err)}")
                analysis_data = {"raw_analysis": analysis_text}
        
        return analysis_data
    
    except requests.exceptions.Timeout:
        logging.error("GPT API呼び出しがタイムアウトしました")
        return {"error": "APIリクエストがタイムアウトしました。しばらく時間をおいて再試行してください。"}
    except requests.exceptions.RequestException as e:
        logging.error(f"GPT API呼び出しネットワークエラー: {str(e)}")
        return {"error": f"ネットワークエラー: {str(e)}"}
    except Exception as e:
        logging.error(f"GPT API呼び出しエラー: {str(e)}")
        return {"error": f"エラーが発生しました: {str(e)}"}

def create_analysis_prompt(resume_text, jd_data):
    """分析用のプロンプトを作成する関数"""
    # JDデータの整形
    # 重要なキーの存在を確認し、存在しない場合は空の文字列を使用
    position = jd_data.get('position', '')
    requirements = jd_data.get('requirements', '')
    responsibilities = jd_data.get('responsibilities', '')
    
    # JDデータのサイズを制限
    position = position[:1000] + ('...' if len(position) > 1000 else '')
    requirements = requirements[:2000] + ('...' if len(requirements) > 2000 else '')
    responsibilities = responsibilities[:2000] + ('...' if len(responsibilities) > 2000 else '')
    
    # 簡略化したJDデータ
    simplified_jd = {
        "position": position,
        "requirements": requirements,
        "responsibilities": responsibilities
    }
    
    jd_json = json.dumps(simplified_jd, ensure_ascii=False, indent=2)
    
    prompt = f"""レジュメ情報:
{resume_text}

JD情報:
{jd_json}

指示：
以下の分析を行い、必ず有効なJSONフォーマットで結果を返してください:

1. このレジュメのスキル・経験と、JDの要件との適合度を100点満点で評価し、その理由を説明してください。

2. この候補者に対して、選択されたポジションに関する面接質問を5個作成してください。質問は候補者の経歴に基づき、技術的能力、経験、文化的フィット、成果などを評価できるものにしてください。

3. 候補者の経歴において掘り下げるべき重要ポイントや、JDとの比較で不足していると思われるスキル・経験を3〜5点特定してください。

4. レジュメに記載された各勤務先とその期間、役職を抽出してリスト化してください。

出力形式:
必ず以下の形式の有効なJSONオブジェクトだけを返してください:
{{
    "適合度評価": {{
        "スコア": 80,
        "評価理由": "..."
    }},
    "面接質問リスト": [
        "質問1...",
        "質問2...",
        "..."
    ],
    "掘り下げポイント": [
        {{"ポイント": "...", "理由": "..."}},
        {{"ポイント": "...", "理由": "..."}},
        "..."
    ],
    "不足スキル": [
        {{"スキル": "...", "重要度": "高/中/低"}},
        {{"スキル": "...", "重要度": "高/中/低"}},
        "..."
    ],
    "勤務履歴": [
        {{"会社名": "...", "役職": "...", "期間": "..."}},
        {{"会社名": "...", "役職": "...", "期間": "..."}},
        "..."
    ]
}}

この形式に厳密に従い、JSONオブジェクトのみを出力してください。文章や説明などJSONの前後に余計なテキストは含めないでください。
"""
    return prompt

def serialize_entity(entity):
    """エンティティをJSONシリアライズ可能な形式に変換"""
    if hasattr(entity, 'dict'):
        return entity.dict
    else:
        result = {}
        for key, value in entity.items():
            if key in ['PartitionKey', 'RowKey', 'Timestamp', 'etag']:
                continue
            result[key] = value
        return result

@app.route(route="ManageJD", methods=["POST", "GET"])
def ManageJD(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('JD管理関数が呼び出されました')
    
    # CORS対応ヘッダー
    headers = {
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Methods": "GET, POST, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type, Authorization"
    }
    
    # OPTIONSリクエストの処理
    if req.method == "OPTIONS":
        return func.HttpResponse(
            status_code=200,
            headers=headers
        )
        
    try:
        if req.method == "GET":
            try:
                category = req.params.get('category')
                jd_id = req.params.get('id')
                
                query = "PartitionKey eq 'jd'"
                if jd_id:
                    query = f"PartitionKey eq 'jd' and RowKey eq '{jd_id}'"
                elif category:
                    query += f" and category eq '{category}'"
                
                table_client = get_table_client("jdList")
                jd_list = list(table_client.query_entities(query))
                
                result = []
                for entity in jd_list:
                    try:
                        jd_data = json.loads(entity.get('data', '{}'))
                    except:
                        jd_data = {}
                    
                    result.append({
                        'id': entity['RowKey'],
                        'title': entity.get('title', ''),
                        'category': entity.get('category', ''),
                        'createdAt': entity.get('createdAt', ''),
                        'updatedAt': entity.get('updatedAt', ''),
                        'data': jd_data
                    })
                
                return func.HttpResponse(
                    json.dumps(result),
                    mimetype="application/json",
                    headers=headers,
                    status_code=200
                )
            except Exception as e:
                logging.error(f"JD一覧取得エラー: {str(e)}")
                return func.HttpResponse(
                    json.dumps({"error": str(e)}),
                    mimetype="application/json",
                    headers=headers,
                    status_code=500
                )
        
        elif req.method == "POST":
            try:
                req_body = req.get_json()
            except ValueError as e:
                logging.error(f"JSONパースエラー: {str(e)}")
                body_text = req.get_body().decode('utf-8', errors='replace')
                logging.error(f"受信したリクエストボディ: {body_text[:1000]}")
                return func.HttpResponse(
                    json.dumps({"error": "Invalid JSON data"}),
                    mimetype="application/json",
                    headers=headers,
                    status_code=400
                )
            
            try:
                jd_id = req_body.get('id')
                is_delete = req_body.get('delete', False)
                
                if is_delete and jd_id:
                    table_client = get_table_client("jdList")
                    table_client.delete_entity(partition_key='jd', row_key=jd_id)
                    return func.HttpResponse(
                        json.dumps({"status": "success", "message": "JDが削除されました"}),
                        mimetype="application/json",
                        headers=headers,
                        status_code=200
                    )
                
                jd_data = req_body.get('data', {})
                
                if not jd_id:
                    jd_id = str(uuid.uuid4())
                
                current_time = datetime.utcnow().isoformat()
                
                table_client = get_table_client("jdList")
                entity = {
                    'PartitionKey': 'jd',
                    'RowKey': jd_id,
                    'title': jd_data.get('position', ''),
                    'category': jd_data.get('category', ''),
                    'updatedAt': current_time,
                    'data': json.dumps(jd_data)
                }
                
                try:
                    existing = table_client.get_entity('jd', jd_id)
                    entity['createdAt'] = existing.get('createdAt', current_time)
                except:
                    entity['createdAt'] = current_time
                
                table_client.upsert_entity(entity)
                
                return func.HttpResponse(
                    json.dumps({"id": jd_id, "status": "success"}),
                    mimetype="application/json",
                    headers=headers,
                    status_code=200
                )
            except Exception as e:
                logging.error(f"JD登録/更新エラー: {str(e)}")
                return func.HttpResponse(
                    json.dumps({"error": str(e)}),
                    mimetype="application/json",
                    headers=headers,
                    status_code=500
                )
    
    except Exception as e:
        logging.error(f"予期せぬエラー: {str(e)}")
        return func.HttpResponse(
            json.dumps({"error": str(e)}),
            mimetype="application/json",
            headers=headers,
            status_code=500
        )

@app.route(route="ImportJDFromExcel", methods=["POST", "OPTIONS"])
def ImportJDFromExcel(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Excel JDインポート関数が呼び出されました')
    
    # CORS対応ヘッダー
    headers = {
        "Access-Control-Allow-Origin": "*",
        "Access-Control-Allow-Methods": "POST, OPTIONS",
        "Access-Control-Allow-Headers": "Content-Type, Authorization"
    }
    
    # OPTIONSリクエストの処理
    if req.method == "OPTIONS":
        return func.HttpResponse(
            status_code=200,
            headers=headers
        )
        
    try:
        file_data = req.get_body()
        
        jd_list = parse_jd_excel(file_data)
        
        results = []
        table_client = get_table_client("jdList")
        
        current_time = datetime.utcnow().isoformat()
        
        for jd in jd_list:
            jd_id = str(uuid.uuid4())
            
            entity = {
                'PartitionKey': 'jd',
                'RowKey': jd_id,
                'title': jd.get('position', ''),
                'category': jd.get('category', ''),
                'createdAt': current_time,
                'updatedAt': current_time,
                'data': json.dumps(jd)
            }
            
            table_client.upsert_entity(entity)
            results.append({
                "id": jd_id, 
                "title": jd.get('position', '')
            })
        
        return func.HttpResponse(
            json.dumps({
                "status": "success", 
                "imported": len(results), 
                "items": results
            }),
            mimetype="application/json",
            headers=headers,
            status_code=200
        )
    except Exception as e:
        logging.error(f"